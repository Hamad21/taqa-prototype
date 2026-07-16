# -*- coding: utf-8 -*-
"""Generate the IT handover / deployment-request document (.docx)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY="00093F"; TEAL="0AB3A1"; LIGHT="EEF1F6"; GREY="5B5D70"
OUT = os.path.join(os.path.dirname(__file__), "..", "Quality-Portal-Handover.docx")

def shade(cell, c):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),c); tcPr.append(shd)
def borders(table, color="C9CCDA"):
    b = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{e}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:space'),'0'); el.set(qn('w:color'),color); b.append(el)
    table._tbl.tblPr.append(b)
def widths(t, ws):
    t.autofit=False
    for row in t.rows:
        for i,w in enumerate(ws): row.cells[i].width=Cm(w)
def crun(cell, text, bold=False, size=9.5, color="1A1A1D", align=None):
    p=cell.paragraphs[0]
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"; r.font.color.rgb=RGBColor.from_string(color); return p
def para(doc, text, bold=False, size=10.5, color="1A1A1D", after=6, before=0, italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size); r.font.name="Arial"; r.font.color.rgb=RGBColor.from_string(color); return p
def h(doc, n, text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"{n}.  {text}"); r.bold=True; r.font.size=Pt(12.5); r.font.name="Arial"; r.font.color.rgb=RGBColor.from_string(NAVY); return p
def bullet(doc, text, boldlead=None):
    p=doc.add_paragraph(style=None); p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.space_after=Pt(3)
    r0=p.add_run("•  "); r0.font.size=Pt(10.5); r0.font.name="Arial"; r0.font.color.rgb=RGBColor.from_string(TEAL); r0.bold=True
    if boldlead:
        rb=p.add_run(boldlead); rb.bold=True; rb.font.size=Pt(10.5); rb.font.name="Arial"; rb.font.color.rgb=RGBColor.from_string("1A1A1D")
    r=p.add_run(text); r.font.size=Pt(10.5); r.font.name="Arial"; r.font.color.rgb=RGBColor.from_string("1A1A1D"); return p

doc=Document()
for s in doc.sections:
    s.top_margin=Mm(16); s.bottom_margin=Mm(16); s.left_margin=Mm(18); s.right_margin=Mm(18)

# Title band
t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; widths(t,[12.9,4.5])
L,R=t.rows[0].cells; shade(L,NAVY); shade(R,TEAL)
lp=L.paragraphs[0]; lp.paragraph_format.space_before=Pt(6); lp.paragraph_format.space_after=Pt(2)
r1=lp.add_run("POWER QUALITY TEAM  ·  POWER TECHNICAL SUPPORT\n"); r1.bold=True; r1.font.size=Pt(8); r1.font.name="Arial"; r1.font.color.rgb=RGBColor.from_string("9AD071")
r2=lp.add_run("Quality Team Portal"); r2.bold=True; r2.font.size=Pt(15); r2.font.name="Arial"; r2.font.color.rgb=RGBColor.from_string("FFFFFF")
lp2=L.add_paragraph(); lp2.paragraph_format.space_before=Pt(1); lp2.paragraph_format.space_after=Pt(6)
r3=lp2.add_run("IT Deployment / Handover Request"); r3.italic=True; r3.font.size=Pt(10); r3.font.name="Arial"; r3.font.color.rgb=RGBColor.from_string("C3C9E6")
rp=R.paragraphs[0]; rp.alignment=WD_ALIGN_PARAGRAPH.CENTER; rp.paragraph_format.space_before=Pt(12)
rr=rp.add_run("TAQA\nDISTRIBUTION"); rr.bold=True; rr.font.size=Pt(12); rr.font.name="Arial"; rr.font.color.rgb=RGBColor.from_string("FFFFFF")

# meta strip
m=doc.add_table(rows=2, cols=4); borders(m); widths(m,[4.35,4.35,4.35,4.35])
for i,(k,v) in enumerate([("Prepared for","IT Department"),("Prepared by","______________________"),("Team","Power Quality Team"),("Date","______________")]):
    shade(m.rows[0].cells[i], LIGHT); crun(m.rows[0].cells[i], k.upper(), bold=True, size=7.5, color=GREY); crun(m.rows[1].cells[i], v, size=9.5)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

h(doc,1,"Purpose")
para(doc,"The Power Quality Team (Power Technical Support) has prepared an internal web portal that gives field and office teams a single place to find controlled quality documents and technical guidance. This document asks IT to host the portal on an approved internal location and to advise on the preferred platform.")

h(doc,2,"What the portal contains")
bullet(doc,"a landing page with inspection checklists, a Quality Awareness Corner, and a contacts section;")
bullet(doc,"a dedicated “Jointing in Weather Conditions” guidance page; and")
bullet(doc,"nine downloadable documents: six controlled checklists (Word) plus three supporting awareness PDFs, each with a reference number and revision (register below).")
reg=doc.add_table(rows=1, cols=4); borders(reg); widths(reg,[3.6,9.0,2.2,2.6])
for i,(txt,al) in enumerate([("Reference",None),("Document",None),("Type",WD_ALIGN_PARAGRAPH.CENTER),("Last updated",WD_ALIGN_PARAGRAPH.CENTER)]):
    shade(reg.rows[0].cells[i], NAVY); crun(reg.rows[0].cells[i], txt, bold=True, size=8.5, color="FFFFFF", align=al)
rows=[("PP.PPS.PTS.QT.04","Pre-execution & Trench Excavation Checklist","Word","Rev 1 · Mar 2025"),
      ("PP.PPS.PTS.QT.05","Cable Laying Checklist","Word","Rev 1 · Mar 2025"),
      ("PP.PPS.PTS.QT.06","Cable Joint Checklist","Word","Rev 1 · Mar 2025"),
      ("PP.PPS.PTS.QT.06A","Cable Joint Guide: Uploading Pictures (Maximo)","Word","Rev 1 · Mar 2025"),
      ("PP.PPS.PTS.QT.07","Equipment Installation Checklist","Word","Rev 1 · Mar 2025"),
      ("PP.PPS.PTS.QT.08","Cable Termination Checklist","Word","Rev 1 · Mar 2025"),
      ("QT-PTS-QA-001","Common Jointing Defects & How to Avoid Them (awareness)","PDF","05 Jul 2026"),
      ("QT-PTS-QA-002","Workmanship Standards for Cable Accessories (awareness)","PDF","28 Jun 2026"),
      ("QT-PTS-QA-003","Jointing in Weather Conditions: Field Guide (awareness)","PDF","12 Jul 2026")]
for ref,name,typ,dt in rows:
    c=reg.add_row().cells
    crun(c[0],ref,size=9,color="6B42D1"); crun(c[1],name,size=9.5); crun(c[2],typ,size=9,align=WD_ALIGN_PARAGRAPH.CENTER); crun(c[3],dt,size=9,align=WD_ALIGN_PARAGRAPH.CENTER,color=GREY)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

h(doc,3,"How it is built (for IT)")
bullet(doc,"a set of HTML, CSS and JavaScript files plus the document downloads.", boldlead="Static website: ")
bullet(doc,"there is no server-side code, no database, no user login and no data is collected or submitted.", boldlead="No backend: ")
bullet(doc,"the pages call Google Fonts (fonts.googleapis.com); this can be self-hosted on request so the site makes no external calls.", boldlead="One external dependency: ")
bullet(doc,"it can be served from any standard web server, static-hosting service, or file share; total size is a few megabytes.", boldlead="Portable: ")
bullet(doc,"the attached ZIP contains the complete site; open index.html to preview it locally.", boldlead="Self-contained: ")

h(doc,4,"Suggested hosting options")
para(doc,"Any of the following would work; IT to advise which fits our environment and policy:", after=4)
bullet(doc,"host the pages and documents on a SharePoint Online site / document library.", boldlead="Microsoft 365 / SharePoint: ")
bullet(doc,"publish the folder on an internal IIS or web server behind the corporate network.", boldlead="Internal web server: ")
bullet(doc,"deploy as an Azure Static Web App with access restricted to the organisation.", boldlead="Azure Static Web Apps: ")
bullet(doc,"embed or link it from the existing intranet.", boldlead="Intranet: ")

h(doc,5,"Security & data handling")
bullet(doc,"The portal is informational only: it displays content and offers document downloads.")
bullet(doc,"No forms post anywhere; no personal data is captured, stored or transmitted.")
bullet(doc,"Access control (if required) can be handled entirely by the chosen hosting platform.")

h(doc,6,"What we are requesting from IT")
bullet(doc,"host the portal at an approved internal location and provide the internal URL;")
bullet(doc,"advise the preferred platform from section 4 (or an alternative that meets policy); and")
bullet(doc,"confirm any change-control or review steps we should follow for future updates.")

h(doc,7,"Points to confirm before go-live")
bullet(doc,"there is a temporary public preview on GitHub Pages used only to demonstrate the portal; it should be replaced by the internal version and can then be taken down.", boldlead="Move off the public preview: ")
bullet(doc,"the portal uses TAQA Distribution branding and should be approved through the usual brand / communications channel.", boldlead="Branding sign-off: ")
bullet(doc,"the checklist and awareness content is Power Quality Team good-practice; confirm ownership and where the controlled master copies will live.", boldlead="Document ownership: ")
bullet(doc,"the portal lists the Power Quality Team contacts (team inbox powerquality.team@taqadistribution.com); please verify all addresses before go-live.", boldlead="Contact details: ")

para(doc," ", after=4)
para(doc,"Attached: Quality-Portal.zip (complete website and documents) and a live preview link provided separately.", italic=True, size=9.5, color=GREY)

# footer
fp=doc.sections[0].footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run("Power Quality Team, Power Technical Support  ·  TAQA Distribution  ·  Internal deployment request")
fr.font.size=Pt(7.5); fr.font.name="Arial"; fr.font.color.rgb=RGBColor.from_string(GREY)

doc.save(os.path.abspath(OUT))
print("saved", os.path.abspath(OUT))
