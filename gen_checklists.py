# -*- coding: utf-8 -*-
"""Generate branded Quality Team inspection checklists (.docx)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = "00093F"
TEAL = "0AB3A1"
LIGHT = "EEF1F6"
GREYTXT = "5B5D70"
OUT = os.path.join(os.path.dirname(__file__), "documents", "checklists")
os.makedirs(OUT, exist_ok=True)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_borders(table, color="C9CCDA", sz="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz); e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)

def set_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)

def run(cell, text, bold=False, size=9.5, color="1A1A1D", align=None, italic=False):
    p = cell.paragraphs[0]
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string(color)
    return p

def para(doc, text, bold=False, size=10.5, color="1A1A1D", after=6, before=0, align=None, upper=False, ls=False, italic=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text.upper() if upper else text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = "Arial"; r.font.color.rgb = RGBColor.from_string(color)
    return p

def build(meta, sections, fname):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Mm(15); s.bottom_margin = Mm(15); s.left_margin = Mm(16); s.right_margin = Mm(16)

    # ---- Title band ----
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_widths(t, [13.2, 4.6])
    left, right = t.rows[0].cells
    shade(left, NAVY); shade(right, TEAL)
    lp = left.paragraphs[0]; lp.paragraph_format.space_after = Pt(2); lp.paragraph_format.space_before = Pt(4)
    r1 = lp.add_run("QUALITY TEAM  ·  POWER TECHNICAL SUPPORT\n"); r1.bold = True; r1.font.size = Pt(8); r1.font.name="Arial"; r1.font.color.rgb = RGBColor.from_string("9AD071")
    r2 = lp.add_run(meta["title"]); r2.bold = True; r2.font.size = Pt(13); r2.font.name="Arial"; r2.font.color.rgb = RGBColor.from_string("FFFFFF")
    lp2 = left.add_paragraph(); lp2.paragraph_format.space_before = Pt(2); lp2.paragraph_format.space_after = Pt(4)
    r3 = lp2.add_run("Inspection & Quality Checklist"); r3.font.size = Pt(9); r3.italic=True; r3.font.name="Arial"; r3.font.color.rgb = RGBColor.from_string("C3C9E6")
    rp = right.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER; rp.paragraph_format.space_before = Pt(10)
    rr = rp.add_run("TAQA\nDISTRIBUTION"); rr.bold=True; rr.font.size=Pt(12); rr.font.name="Arial"; rr.font.color.rgb=RGBColor.from_string("FFFFFF")

    # ---- Meta strip ----
    m = doc.add_table(rows=2, cols=4); set_borders(m); set_widths(m, [4.45,4.45,4.45,4.45])
    labels = [("Reference No.", meta["ref"]), ("Revision", meta["rev"]),
              ("Last Updated", meta["date"]), ("Discipline", meta["discipline"])]
    for i,(k,v) in enumerate(labels):
        top = m.rows[0].cells[i]; bot = m.rows[1].cells[i]
        shade(top, LIGHT)
        run(top, k.upper(), bold=True, size=7.5, color=GREYTXT)
        run(bot, v, bold=(i==0), size=9.5, color=(NAVY if i==0 else "1A1A1D"))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- Project fill-in strip ----
    pj = doc.add_table(rows=2, cols=4); set_borders(pj); set_widths(pj,[4.45,4.45,4.45,4.45])
    for i,k in enumerate(["Project / Contract","Location / Feeder","Contractor","Date of Inspection"]):
        shade(pj.rows[0].cells[i], "FBFBFD")
        run(pj.rows[0].cells[i], k.upper(), bold=True, size=7.5, color=GREYTXT)
        run(pj.rows[1].cells[i], " ", size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- Sections ----
    for si,(sname, items) in enumerate(sections, 1):
        para(doc, f"{si}.  {sname}", bold=True, size=11, color=NAVY, after=4, before=6, upper=True)
        tb = doc.add_table(rows=1, cols=5); set_borders(tb); set_widths(tb,[1.1,10.2,1.5,1.5,3.5])
        hdr = tb.rows[0].cells
        for c,txt,al in [(0,"#",WD_ALIGN_PARAGRAPH.CENTER),(1,"Check item",None),(2,"Yes/No",WD_ALIGN_PARAGRAPH.CENTER),(3,"N/A",WD_ALIGN_PARAGRAPH.CENTER),(4,"Remarks",None)]:
            shade(hdr[c], NAVY); run(hdr[c], txt, bold=True, size=8.5, color="FFFFFF", align=al)
        for j,item in enumerate(items,1):
            cells = tb.add_row().cells
            run(cells[0], str(j), size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREYTXT)
            run(cells[1], item, size=9.5)
            run(cells[2], "", align=WD_ALIGN_PARAGRAPH.CENTER)
            run(cells[3], "", align=WD_ALIGN_PARAGRAPH.CENTER)
            run(cells[4], "")
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- Legend ----
    para(doc, "Mark each item Yes (compliant), No (non-conformance — raise on the remarks line) or N/A. "
              "Any “No” must be closed out before the activity is accepted.", italic=True, size=8.5, color=GREYTXT, before=2, after=10)

    # ---- Sign-off ----
    para(doc, "Sign-off", bold=True, size=11, color=NAVY, upper=True, after=4)
    so = doc.add_table(rows=2, cols=3); set_borders(so); set_widths(so,[5.93,5.93,5.94])
    for i,k in enumerate(["Inspected by (name / signature)","Reviewed by — Quality Team","Date"]):
        shade(so.rows[0].cells[i], LIGHT); run(so.rows[0].cells[i], k.upper(), bold=True, size=7.5, color=GREYTXT)
        run(so.rows[1].cells[i], " ", size=9.5)
        so.rows[1].cells[i].paragraphs[0].paragraph_format.space_before = Pt(10)

    # ---- Footer ----
    foot = doc.sections[0].footer
    fp = foot.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f'{meta["ref"]}  ·  Rev {meta["rev"]}  ·  Last updated {meta["date"]}  ·  '
                    f'Quality Team, Power Technical Support  ·  Controlled document — verify current revision before use')
    fr.font.size = Pt(7.5); fr.font.name="Arial"; fr.font.color.rgb = RGBColor.from_string(GREYTXT)

    path = os.path.join(OUT, fname); doc.save(path)
    print("saved", os.path.basename(path))

# ---------------------------------------------------------------- catalogue
build(
  dict(title="LV & MV Cable Jointing Inspection Checklist", ref="QT-PTS-CHK-001", rev="2.0",
       date="08 Jul 2026", discipline="Power — Cable Jointing"),
  [
    ("Pre-jointing preparation", [
        "Correct joint kit selected for cable type, size and voltage class",
        "All materials within shelf life and undamaged; packaging intact until use",
        "Jointing enclosure / tent erected; area clean, dry and protected from weather",
        "Cable ends undamaged; correct overlap and working length available",
        "Correct tooling, crimp tool and dies available and calibrated where required",
        "Approved drawing / manufacturer instruction available at the work face",
        "PPE and safe isolation confirmed before work begins",
    ]),
    ("Jointing", [
        "Cable preparation dimensions set out per manufacturer instruction",
        "Insulation and semi-con screen cut cleanly, no nicks to conductor",
        "Conductor connector correct size; crimped with correct die and number of crimps",
        "Connector alignment and inter-phase spacing correct",
        "Insulation / stress control applied without voids, trapped air or contamination",
        "Metallic screen continuity maintained and bonded as specified",
        "Earth continuity conductor connected and continuous across the joint",
    ]),
    ("Post-jointing & records", [
        "Joint protection / mechanical shell installed correctly",
        "Joint position, phasing and identification recorded",
        "Photographic record taken at key stages",
        "Backfill / protection reinstated as specified",
        "Work area cleared; waste and offcuts removed",
    ]),
  ],
  "QT-PTS-CHK-001_Cable-Jointing-Inspection.docx")

build(
  dict(title="Cable Laying & Backfilling Inspection Checklist", ref="QT-PTS-CHK-002", rev="1.3",
       date="24 Jun 2026", discipline="Power — Cable Installation"),
  [
    ("Trench & route", [
        "Trench route as per approved drawing and permit",
        "Trench depth and width to specification",
        "Trench free of sharp objects, debris and standing water",
        "Sand / bedding layer placed to specified thickness",
    ]),
    ("Cable pulling", [
        "Minimum bending radius maintained throughout the pull",
        "Cable rollers / bell-mouth used; no dragging over trench edge",
        "Pulling tension and speed controlled within limits",
        "Cable sheath inspected — no cuts, abrasions or flat spots",
        "Circuit spacing and phase arrangement as specified",
    ]),
    ("Protection & separation", [
        "Warning tape and/or protection tiles at specified depth",
        "Separation from other cables and services maintained",
        "Route markers / duct references installed where required",
    ]),
    ("Backfill & reinstatement", [
        "Surround material approved and free of stones near the cable",
        "Backfill compacted in layers to specification",
        "As-laid measurements and route recorded before covering",
        "Surface reinstated and site left tidy",
    ]),
  ],
  "QT-PTS-CHK-002_Cable-Laying-Backfilling.docx")

build(
  dict(title="Jointing in Adverse Weather — Pre-Work Checklist", ref="QT-PTS-CHK-003", rev="1.1",
       date="12 Jul 2026", discipline="Power — Cable Jointing (Weather)"),
  [
    ("Conditions recorded", [
        "Ambient temperature measured and recorded",
        "Relative humidity measured and recorded",
        "Weather forecast checked for the planned jointing window",
        "Wind / airborne dust conditions assessed",
    ]),
    ("Controls in place", [
        "Enclosure / tent erected and sealed against wind, rain and dust",
        "Insulation surfaces and materials clean and dry",
        "Approved warming / dehumidifying means available if needed",
        "Excavation free of standing water; dry working platform provided",
        "Materials within their stated temperature handling window",
    ]),
    ("Go / no-go decision", [
        "No active rain on, or forecast for, the open joint",
        "No active sandstorm; enclosure cleaned down",
        "Surfaces confirmed clean, dry and in temperature range at close-up",
        "Decision to PROCEED / STOP recorded with name and time",
    ]),
  ],
  "QT-PTS-CHK-003_Adverse-Weather-Pre-Work.docx")

build(
  dict(title="Substation Cable Termination Quality Checklist", ref="QT-PTS-CHK-004", rev="1.0",
       date="02 Jul 2026", discipline="Power — Terminations"),
  [
    ("Preparation", [
        "Correct termination kit for cable and switchgear interface",
        "Gland type, size and entry method as specified",
        "Electrical clearances and creepage distances confirmed",
        "Work area clean, dry and adequately lit",
    ]),
    ("Termination", [
        "Lug size correct; crimped with correct die and crimp count",
        "Stress control / boot / heat-shrink applied without voids",
        "Phase identification applied correctly",
        "Connections torqued to specified value and marked",
    ]),
    ("Earthing & testing", [
        "Cable screen / armour earthed and continuous",
        "Gland earthing installed and bonded",
        "Phasing verified against the network",
        "Insulation resistance test carried out and recorded",
    ]),
    ("Records & housekeeping", [
        "Test values and torque records attached",
        "Photographic record taken",
        "Area cleared; covers and barriers reinstated",
    ]),
  ],
  "QT-PTS-CHK-004_Substation-Termination.docx")

print("done")
